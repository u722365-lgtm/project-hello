#!/usr/bin/env python3
"""Generate Word (.docx) for ShadowTalk Podcast Script Episode 1."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "public/downloads/ShadowTalk-Podcast-Script-Episode-1.docx"

SEGMENTS = [
    ("Segment 1 — Cold open (0:00–1:30)", [
        ("FRIEND", 'Okay, real talk — how many of you have typed something into ChatGPT and immediately thought… "wait, should I have said that?"'),
        ("ZAIN", "Literally everyone. You ask about a personal problem, a business idea, something embarrassing — and it's just… there. Saved. Processed. Somewhere in the cloud."),
        ("FRIEND", "And you have no idea who can see it, who trained on it, or if it'll show up again in six months."),
        ("ZAIN", "Exactly. That's why we built ShadowTalk AI — private AI chat for people who actually care about what happens to their words."),
        ("FRIEND", "I'm [Friend's name]."),
        ("ZAIN", "And I'm Zain — founder of ShadowTalk. We're two guys from Pakistan trying to build something the big AI companies won't."),
        ("FRIEND", "Let's get into it."),
    ]),
    ("Segment 2 — Intro & hook (1:30–4:00)", [
        ("FRIEND", "So Zain — before we explain features and all that — why did you start this? You're 17. You could've just… used ChatGPT like everyone else."),
        ("ZAIN", "Honestly? I was using it. Everyone is. But the more I built with AI — code, ideas, personal stuff — the more uncomfortable I got. It felt like every conversation was a permanent record I never agreed to."),
        ("FRIEND", "Like leaving your diary open in a coffee shop."),
        ("ZAIN", "Perfect analogy. And I kept thinking: students use this for homework and personal questions. Developers paste code, API keys, client data. Founders talk about ideas they haven't launched yet. All of that goes somewhere."),
        ("FRIEND", 'And the big companies\' answer is basically "trust us."'),
        ("ZAIN", "Right. ShadowTalk's answer is different: privacy first. Stealth mode. Options to work more locally. You control the vibe — not some black box in California."),
        ("FRIEND", 'Okay but be honest — is this just "another ChatGPT wrapper"? Because TikTok is full of those.'),
        ("ZAIN", "Fair question. A wrapper is a thin skin on someone else's API. ShadowTalk is a full platform — chat, agents, mission control, templates, offline options, admin tools, growth systems… I've been building this for a long time. It's not a weekend project with a landing page."),
        ("FRIEND", "So what's the one sentence pitch?"),
        ("ZAIN", "ShadowTalk is private AI for people who take their data seriously — students, builders, and founders who don't want every thought logged forever."),
    ]),
    ("Segment 3 — The problem (4:00–7:00)", [
        ("FRIEND", "Let's make this super relatable. Give me three types of people who should not be using normal AI the way they do today."),
        ("ZAIN", "Easy. One — Students. You ask about mental health, family stuff, relationship advice, even help with assignments — that's all tied to you. Forever."),
        ("FRIEND", "And when you're 17, 18, 22… you're still figuring life out. You don't want that archived."),
        ("ZAIN", "Two — Developers. You paste code, error logs, sometimes credentials by mistake. One bad paste and you're exposed."),
        ("FRIEND", 'We\'ve all done the "oh no" scroll-back.'),
        ("ZAIN", "Three — Founders and creators. You brainstorm product names, pricing, strategy, things you haven't announced. Why would you train the internet on your unreleased idea?"),
        ("FRIEND", 'So ShadowTalk is for anyone who\'s ever thought: "I wish this conversation didn\'t exist after I closed the tab."'),
        ("ZAIN", "That's literally the brand. Shadow — it fades. Talk — you still get the power of AI."),
    ]),
    ("Segment 4 — What ShadowTalk actually does (7:00–11:00)", [
        ("FRIEND", "Alright, demo time — but podcast demo. Walk me through what someone sees when they open shadowtalk-ai.com."),
        ("ZAIN", "You land on a clean AI chat experience. Sign up free. Start talking. But the difference is in the philosophy and the tools:"),
        ("ZAIN", "First — Privacy & stealth. Stealth mode vibes. You're not here to feed a giant memory machine."),
        ("ZAIN", "Second — Power user stuff. Mission Control, agent tools, templates, themes — you can actually work in it."),
        ("ZAIN", "Third — Offline / local options. For people who want even more control — especially in Pakistan where connectivity and trust are both real issues."),
        ("FRIEND", "What would I use it for as a normal user?"),
        ("ZAIN", "Homework help without feeling watched. Coding help without leaking your project. Planning content, businesses, exams. Even just venting or thinking out loud — safely."),
        ("FRIEND", "And for you as the founder — what's your daily use?"),
        ("ZAIN", "Building ShadowTalk inside ShadowTalk. Debugging. Writing copy. Planning features. Testing agents. I dogfood my own product — if it's bad, I feel it first."),
        ("FRIEND", "What's the feature you're most proud of?"),
        ("ZAIN", "The privacy positioning plus real product depth. A lot of private AI projects are marketing. We actually built the platform — chat, admin, growth engine, video studio, desktop path."),
        ("FRIEND", "Video studio?"),
        ("ZAIN", "Yeah — Pro users can create short-form content inside the app. For a founder in 2026, that matters. You need TikTok, Reels, Shorts."),
    ]),
    ("Segment 5 — Founder story (11:00–14:00)", [
        ("FRIEND", "Let's talk about you for a second. 17. Karachi. Building an AI company. That's not normal."),
        ("ZAIN", "It's not normal, but it's real. I started because I was obsessed with AI and privacy at the same time. Pakistan doesn't get talked about much in global tech — but we have smart people here who deserve world-class tools."),
        ("FRIEND", "What's the hardest part nobody sees?"),
        ("ZAIN", "Two things. One — payments. Stripe basically doesn't work for a 17-year-old founder in Pakistan. So we built JazzCash, Easypaisa, bank transfer, USDT. Two — deployment. You can build forever, but if users don't see the update, it's like you built nothing."),
        ("FRIEND", "So you've had traffic but no revenue?"),
        ("ZAIN", "We've had real visitors — over a thousand in 90 days — people from the US, Pakistan, everywhere. Long session times. People stay. But payments were blocked until recently. That's changing now."),
        ("FRIEND", "What would you tell another young founder in Pakistan listening to this?"),
        ("ZAIN", "Ship. Talk to users. Don't wait for perfect. And solve your problem first — I built ShadowTalk because I wanted it. If you feel the pain, you'll survive the hard months."),
    ]),
    ("Segment 6 — Pricing & how to support (14:00–16:30)", [
        ("FRIEND", "Okay, the question everyone asks: Is it free?"),
        ("ZAIN", "Yes — there's a free tier to try it. When you want more — unlimited messages, premium features, video studio — we have paid plans."),
        ("FRIEND", "Break down the plans simply."),
        ("ZAIN", "Pro — Rs 1,499/month — great for students. Premium — Rs 3,999/month — power users. Elite — Rs 5,999/month — founders and teams. International users can pay in USD or USDT too."),
        ("FRIEND", "And Pakistanis can pay how?"),
        ("ZAIN", "JazzCash and Easypaisa — number 03211798561 — or bank transfer to Meezan Bank. You pay, upload your receipt on the site, we verify within 24 hours — usually much faster."),
        ("FRIEND", "Why not just use Stripe like every SaaS?"),
        ("ZAIN", "Because I'm in Pakistan and seventeen. The system wasn't built for me. So we built our own path instead of waiting for permission."),
        ("FRIEND", "Respect."),
    ]),
    ("Segment 7 — Hot takes / debate (16:30–19:00)", [
        ("FRIEND", 'Hot take 1: "If you\'re not paying for AI, you\'re the product."'),
        ("ZAIN", "Mostly true. Free tiers are acquisition. Your data, your attention, your behavior — that's the currency. ShadowTalk's model is: free to try, pay for real value."),
        ("FRIEND", 'Hot take 2: "Privacy AI is for criminals."'),
        ("ZAIN", "That's what they said about encrypted messaging too. Privacy is for everyone — journalists, students, doctors, founders, normal people who text things they don't want leaked."),
        ("FRIEND", 'Hot take 3: "Pakistan can\'t build global AI products."'),
        ("ZAIN", "Watch us. We already have nearly half our traffic from the US. The internet doesn't care where your passport is if your product solves a real problem."),
    ]),
    ("Segment 8 — Closing & call to action (19:00–21:00)", [
        ("FRIEND", "We're wrapping up. If someone listened to this whole thing — what should they do right now?"),
        ("ZAIN", "Three steps: One — go to shadowtalk-ai.com and try it free. Two — if you love it, upgrade to Pro at shadowtalk-ai.com/founder-access?plan=pro. Three — share this podcast with one friend who still pastes sensitive stuff into ChatGPT without thinking."),
        ("FRIEND", "Where can people follow you?"),
        ("ZAIN", "Website: shadowtalk-ai.com. LinkedIn: Zain Ahmed — ShadowTalk AI. TikTok: we're posting clips from this conversation."),
        ("FRIEND", "And if you're a student in Pakistan listening — JazzCash Pro is literally less than some monthly phone packages. Private AI is not just for Silicon Valley anymore."),
        ("ZAIN", "That's the mission. AI for everyone — but your data stays yours."),
        ("FRIEND", "This was Episode 1 of ShadowTalk."),
        ("ZAIN", "Episode 2 — we'll do a live demo and roast my own product bugs. Stay tuned."),
        ("BOTH", "Peace."),
    ]),
]


def add_line(doc, speaker, text):
    p = doc.add_paragraph()
    run = p.add_run(f"{speaker}: ")
    run.bold = True
    if speaker == "ZAIN":
        run.font.color.rgb = RGBColor(67, 56, 202)
    elif speaker == "FRIEND":
        run.font.color.rgb = RGBColor(5, 150, 105)
    else:
        run.font.color.rgb = RGBColor(107, 114, 128)
    p.add_run(text)


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("ShadowTalk AI — Podcast Script", 0)
    doc.add_paragraph("Episode 1: Why we built AI that forgets on purpose")
    doc.add_paragraph("Hosts: Zain Ahmed (Founder) + Co-host (Friend)")
    doc.add_paragraph("Length: ~20 minutes · June 2026")
    doc.add_paragraph("Website: https://www.shadowtalk-ai.com")
    doc.add_page_break()

    doc.add_heading("Pre-show checklist", 1)
    for item in [
        "Mic check — both hosts at equal volume",
        "Water nearby, phone on silent",
        "URL on screen: shadowtalk-ai.com",
        "JazzCash number ready: 03211798561",
        "Vibe off the script — don't read robotically",
    ]:
        doc.add_paragraph(f"☐ {item}", style="List Bullet")

    doc.add_page_break()

    for title, lines in SEGMENTS:
        doc.add_heading(title, 1)
        for speaker, text in lines:
            add_line(doc, speaker, text)
        doc.add_paragraph()

    doc.add_page_break()
    doc.add_heading("Bonus — TikTok / Reels clips", 1)
    clips = [
        "ChatGPT remembers everything. We don't. — Zain",
        "I'm 17 and Stripe said no — so we use JazzCash. — Zain",
        "Three people who shouldn't use normal AI — Friend asks, Zain answers",
        "Rs 1,499 for private AI — cheaper than your Netflix — Friend",
        "Hot take: privacy isn't for criminals — Zain",
        "We get 48% traffic from the US — built in Karachi — Zain",
    ]
    for c in clips:
        doc.add_paragraph(c, style="List Number")

    doc.add_heading("Key facts quick reference", 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Topic"
    table.rows[0].cells[1].text = "Detail"
    facts = [
        ("Pro price", "Rs 1,499/month"),
        ("Premium", "Rs 3,999/month"),
        ("Elite", "Rs 5,999/month"),
        ("JazzCash / Easypaisa", "03211798561"),
        ("Bank IBAN", "PK08 MEZN 0099 1701 1274 9131"),
        ("Checkout URL", "shadowtalk-ai.com/founder-access?plan=pro"),
        ("Traffic", "~1,100 visitors / 90 days, 48% US, 20% PK"),
    ]
    for k, v in facts:
        row = table.add_row().cells
        row[0].text = k
        row[1].text = v

    doc.add_paragraph()
    footer = doc.add_paragraph(
        "ShadowTalk AI — Podcast Script Episode 1 · June 2026 · Prepared for Zain Ahmed & co-host"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in footer.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(156, 163, 175)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
