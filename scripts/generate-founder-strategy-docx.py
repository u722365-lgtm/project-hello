#!/usr/bin/env python3
"""Generate Word (.docx) version of the Founder Strategy Brief."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "public/downloads/ShadowTalk-Founder-Strategy-2026.docx"


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Cover
    p = doc.add_paragraph("Confidential — Founder Strategy Brief")
    p.runs[0].font.color.rgb = RGBColor(99, 102, 241)
    p.runs[0].bold = True

    title = doc.add_heading("ShadowTalk AI", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph(
        "From zero revenue to first paying customers — a practical playbook for Zain Ahmed"
    )
    doc.add_paragraph("Prepared by: AI co-founder / strategy partner (Cursor Cloud Agent)")
    doc.add_paragraph("Founder: Zain Ahmed — Founder & Lead Architect, Karachi, Pakistan")
    doc.add_paragraph("Date: June 5, 2026")
    doc.add_paragraph("Live site: https://www.shadowtalk-ai.com")
    doc.add_paragraph("LinkedIn: linkedin.com/in/zain-ahmed-917b6b3a6")
    doc.add_paragraph("Repo: zain836/shadowtalk-ai-903ca615")
    doc.add_page_break()

    doc.add_heading("1. Executive Summary", 1)
    doc.add_paragraph(
        "ShadowTalk AI is a privacy-first AI chat platform built by a 17-year-old founder in Pakistan. "
        "The product has real traction — roughly 1,100 visitors and 3,100 page views over 90 days, "
        "with strong engagement (11+ minute average session) — but $0 revenue because payments were never fully unblocked."
    )
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ("Traffic from United States", "~48%"),
            ("Traffic from Pakistan", "~20%"),
            ("Bounce rate", "63%"),
            ("Pricing page views", "36"),
        ],
    )
    insight = doc.add_paragraph()
    insight.add_run("Core insight: ").bold = True
    insight.add_run(
        "You do not need Stripe from Pakistan to make money today. JazzCash, Easypaisa, bank transfer, "
        "and USDT are live paths. Card checkout via Lemon Squeezy is optional for US users."
    )

    doc.add_heading("2. Founder Context", 1)
    add_table(
        doc,
        ["Factor", "Reality", "Implication"],
        [
            ("Age & location", "17, Pakistan", "Stripe/card merchant accounts are difficult or unavailable"),
            ("Audience split", "US 48%, PK 20%", "Need two payment rails: local wallets + international"),
            ("Product maturity", "Large codebase", "Problem is conversion & deploy, not more features"),
            ("LinkedIn", "480 connections", "Trust channel for first PK customers + B2B leads"),
            ("Revenue", "$0 to date", "Focus on checkout + outreach over new features"),
        ],
    )

    doc.add_heading("3. Product & Platform (What Exists)", 1)
    doc.add_heading("Core product", 2)
    for item in [
        "Private AI chat with stealth mode, offline/local AI, multi-provider support",
        "Mission Control, agent tools, templates, themes, video studio (Pro+)",
        "Full admin panel, growth analytics, referral system",
        "PWA + desktop paths; professional UI on Lovable deployment",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Growth engine — ShadowScale (built)", 2)
    for item in [
        "Database migration + orchestrator/worker edge functions",
        "Admin Growth Command panel at /admin",
        "Referral notifications, share signals, video studio promo, cron automation",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Content & viral assets (built)", 2)
    for item in [
        "Remotion package: 60s 9:16 viral short compositions with TTS voiceover",
        "In-app Shadow Video Studio at /video-studio (Pro+ gated)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4. The Payment Problem (and the Fix)", 1)
    doc.add_heading("Why revenue was $0", 2)
    reasons = [
        "Stripe does not work for a 17-year-old founder in Pakistan",
        "Receipt flow was WhatsApp-only",
        "Admin verify did not activate user plans",
        "Migrations and edge functions not deployed to production",
        "Frontend changes not redeployed on Lovable",
    ]
    for i, r in enumerate(reasons, 1):
        doc.add_paragraph(f"{i}. {r}")

    doc.add_heading("Payment unblock — implemented (PR #105)", 2)
    add_table(
        doc,
        ["Channel", "Details", "Price (Pro)"],
        [
            ("JazzCash", "03211798561", "Rs 1,499/mo"),
            ("Easypaisa", "03211798561", "Rs 1,499/mo"),
            ("Meezan Bank", "PK08 MEZN 0099 1701 1274 9131", "Rs 1,499/mo"),
            ("USDT (TRC20)", "TKfKJ7ESFcnMTd2F1DkrvZ4buCWneAmHqz", "~$5/mo"),
            ("Card (optional)", "Lemon Squeezy", "$5–20/mo"),
        ],
    )

    doc.add_heading("Deploy checklist (production)", 2)
    for item in [
        "Merge PR #105 and redeploy on Lovable",
        "Run migration 20260611120000_payment_receipts_unblock.sql in ShadowTalk backend",
        "Deploy: backend functions deploy verify-manual-payment",
        "Optional: set LEMONSQUEEZY_* and VITE_LEMONSQUEEZY_VARIANT_* env vars",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_page_break()

    doc.add_heading("5. Pricing Strategy", 1)
    add_table(
        doc,
        ["Plan", "USD", "PKR", "Target buyer"],
        [
            ("Pro", "$5/mo", "Rs 1,499", "Students, daily builders"),
            ("Premium", "$15/mo", "Rs 3,999", "Power users, Mission Control"),
            ("Elite", "$20/mo", "Rs 5,999", "Founders, agencies, teams"),
        ],
    )

    doc.add_heading("6. Go-to-Market Playbook", 1)
    doc.add_heading("Week 1 — First 10 paying users", 2)
    for item in [
        "Deploy payment unblock",
        "Post on LinkedIn — problem → product → JazzCash price → link",
        "DM 20 engaged connections; offer Pro at Rs 1,499",
        "WhatsApp status with checkout screenshot",
        "Verify every receipt within 2 hours",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("LinkedIn post template", 2)
    doc.add_paragraph(
        "Most AI tools remember everything you type. ShadowTalk doesn't.\n\n"
        "I built ShadowTalk AI — private AI chat for students and builders in Pakistan (and globally).\n\n"
        "Pro is Rs 1,499/month via JazzCash or Easypaisa. No Stripe needed.\n\n"
        "Try it: shadowtalk-ai.com/founder-access?plan=pro\n\n"
        "#AI #Pakistan #Startup #Privacy"
    ).italic = True

    doc.add_heading("7. What NOT to Do Right Now", 1)
    for item in [
        "Don't build more features before first payment",
        "Don't chase Stripe from Pakistan",
        "Don't add heavy images/animations to landing",
        "Don't merge 50 PRs at once",
        "Don't ignore /pricing — link checkout in every post",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("8. Technical Debt & Production Blockers", 1)
    add_table(
        doc,
        ["Blocker", "Impact", "Fix"],
        [
            ("Migrations not applied", "Features break silently", "Run SQL in ShadowTalk backend dashboard"),
            ("Edge functions not deployed", "Verify flow fails", "Deploy via ShadowTalk backend CLI"),
            ("Lovable not redeployed", "Users see old UI", "Merge + redeploy gemini branch"),
            ("API_ACCESS_TOKEN missing", "CI skips deploy", "Add GitHub secret"),
            ("Resend domain unverified", "Emails fail", "Verify domain (lower priority)"),
        ],
    )

    doc.add_heading("9. Key URLs & Admin Paths", 1)
    add_table(
        doc,
        ["Purpose", "URL"],
        [
            ("Checkout (Pro)", "/founder-access?plan=pro"),
            ("Checkout (Premium)", "/founder-access?plan=premium"),
            ("Pricing page", "/pricing"),
            ("Admin payments", "/admin → Manual Payments"),
            ("Growth Command", "/admin → Growth Command"),
            ("Video Studio", "/video-studio (Pro+)"),
        ],
    )

    doc.add_heading("10. 30-Day Revenue Targets", 1)
    add_table(
        doc,
        ["Milestone", "Target", "How"],
        [
            ("Day 1–3", "Deploy + 1 test payment", "Pay Rs 1,499 yourself, verify in admin"),
            ("Week 1", "3 paying users", "LinkedIn + WhatsApp + college network"),
            ("Week 2", "10 paying users", "~Rs 15k–40k MRR"),
            ("Month 1", "25 paying users", "Referral loop + 1 viral short"),
        ],
    )

    north = doc.add_paragraph()
    north.add_run("North star: ").bold = True
    north.add_run(
        "First Rs 1,499 hitting your JazzCash from a stranger — not another feature shipped."
    )

    doc.add_heading("11. Conversation Log — Major Decisions", 1)
    decisions = [
        "Accepted Pakistan-first payment rails instead of waiting for Stripe",
        "Built in-app receipt upload + admin verify → auto plan activation",
        "Shipped ShadowScale growth engine",
        "Built Remotion viral video + in-app Video Studio",
        "Identified deploy gap as root cause of changes not showing",
        "LinkedIn profile reviewed for trust + checkout links",
        "PR #105: payment gateway unblock ready to merge",
    ]
    for i, d in enumerate(decisions, 1):
        doc.add_paragraph(f"{i}. {d}")

    doc.add_heading("12. Next Actions (Priority Order)", 1)
    actions = [
        "Merge PR #105 → redeploy Lovable",
        "Run payment migration on production ShadowTalk backend",
        "Deploy verify-manual-payment edge function",
        "Test full flow: pay → submit receipt → admin verify → Pro active",
        "Post LinkedIn announcement with checkout link",
        "Set Lemon Squeezy variants for US card checkout (optional)",
        "Record + post 1 viral short from Video Studio / Remotion",
    ]
    for a in actions:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"☐ {a}")

    doc.add_paragraph()
    footer = doc.add_paragraph(
        "ShadowTalk AI — Founder Strategy Brief · June 2026 · Prepared for Zain Ahmed\n"
        "This document reflects strategy discussions between founder and AI co-founder. "
        "Not financial or legal advice."
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in footer.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(156, 163, 175)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
